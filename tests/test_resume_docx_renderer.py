"""Native Microsoft Word resume rendering tests."""

from __future__ import annotations

import unittest
from io import BytesIO
from zipfile import is_zipfile

try:
    from docx import Document
    from docx.enum.text import WD_TAB_ALIGNMENT
except ImportError:
    Document = None
    WD_TAB_ALIGNMENT = None

from models import CandidateProfile, GeneratedResumeContent, ResumeKnowledgeBase
from services import ResumeDocxRenderer


class ResumeDocxRendererTests(unittest.TestCase):
    @unittest.skipIf(Document is None, "python-docx is not installed")
    def test_renders_editable_professional_word_document(self) -> None:
        candidate = CandidateProfile(
            candidate_id="candidate-1",
            full_name="Example Candidate",
            email="candidate@example.com",
            phone="(555) 123-4567",
            location="Denver, CO",
            linkedin_url="https://www.linkedin.com/in/example-candidate",
            github_url="https://github.com/example-candidate",
            website_url="https://example.dev",
            skills=("Python", "SQL"),
        )
        knowledge = ResumeKnowledgeBase.from_dict(
            {
                "candidate_id": "candidate-1",
                "skills": ["Python", "SQL"],
                "roles": [
                    {
                        "company": "Example Corp",
                        "title": "Data Engineer",
                        "location": "Remote, US",
                        "start_date": "2022-01",
                        "end_date": "Present",
                        "responsibilities": ["Built reliable data pipelines."],
                    }
                ],
                "achievements": [
                    {
                        "category": "Reliability",
                        "description": "Improved platform reliability.",
                    }
                ],
                "education": [
                    {
                        "institution": "Example University",
                        "location": "Denver, CO",
                        "degree": "Bachelor of Science",
                        "field": "Computer Engineering",
                    }
                ],
                "certifications": [
                    {
                        "name": "Cloud Certification",
                        "issued": "2025-01",
                        "status": "Current",
                    }
                ],
            }
        )
        content = GeneratedResumeContent.from_dict(
            {
                "professional_summary": "Builds reliable data platforms.",
                "skills": ["Python", "SQL"],
                "experience": [
                    {
                        "company": "Example Corp",
                        "title": "Data Engineer",
                        "location": "Remote, US",
                        "start_date": "2022-01",
                        "end_date": "Present",
                        "responsibilities": ["Built reliable data pipelines."],
                    }
                ],
                "career_highlights": [
                    {
                        "category": "Reliability",
                        "description": "Improved platform reliability.",
                    }
                ],
                "education": [
                    {
                        "institution": "Example University",
                        "location": "Denver, CO",
                        "degree": "Bachelor of Science",
                        "field": "Computer Engineering",
                        "status": None,
                    }
                ],
                "certifications": [
                    {
                        "name": "Cloud Certification",
                        "issued": "2025-01",
                        "status": "Current",
                    }
                ],
            }
        )
        content.validate_against(knowledge, candidate_skills=candidate.skills)

        rendered = ResumeDocxRenderer().render(
            candidate,
            content,
            target_title="Senior Data Engineer",
        )

        self.assertTrue(is_zipfile(BytesIO(rendered)))
        assert Document is not None
        document = Document(BytesIO(rendered))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("Example Candidate", text)
        self.assertIn("candidate@example.com", text)
        self.assertIn("(555) 123-4567", text)
        self.assertIn("https://www.linkedin.com/in/example-candidate", text)
        self.assertIn("https://github.com/example-candidate", text)
        self.assertIn("https://example.dev", text)
        self.assertIn("Senior Data Engineer — Builds reliable data platforms.", text)
        self.assertIn("PROFESSIONAL EXPERIENCE", text)
        self.assertIn("Remote, US", text)
        self.assertIn("Jan 2022 – Current", text)
        self.assertIn("Built reliable data pipelines.", text)
        self.assertIn("Improved platform reliability.", text)
        self.assertIn("Bachelor of Science in Computer Engineering", text)
        self.assertIn("Cloud Certification", text)
        self.assertIn("Issued January 2025", text)
        self.assertEqual(
            document.core_properties.title,
            "Example Candidate - Senior Data Engineer Resume",
        )
        bold_runs = [
            run.text
            for paragraph in document.paragraphs
            for run in paragraph.runs
            if run.bold is True
        ]
        self.assertEqual(bold_runs, ["Data Engineer"])
        self.assertFalse(document.styles["Title"].font.bold)
        self.assertFalse(document.styles["Heading 1"].font.bold)
        experience_heading = next(
            paragraph
            for paragraph in document.paragraphs
            if paragraph.text.startswith("Data Engineer")
        )
        self.assertEqual(
            experience_heading.text.splitlines()[0],
            "Data Engineer\tJan 2022 – Current",
        )
        tab_stops = list(experience_heading.paragraph_format.tab_stops)
        self.assertEqual(len(tab_stops), 1)
        self.assertEqual(tab_stops[0].alignment, WD_TAB_ALIGNMENT.RIGHT)
        self.assertAlmostEqual(tab_stops[0].position.inches, 7.2, places=2)
        self.assertAlmostEqual(document.sections[0].top_margin.inches, 0.55, places=2)


if __name__ == "__main__":
    unittest.main()
