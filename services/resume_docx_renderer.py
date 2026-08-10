"""Render validated resume content as an editable Microsoft Word document."""

from __future__ import annotations

from io import BytesIO
from typing import Any

from models import CandidateProfile, GeneratedResumeContent, GeneratedResumeRole
from utils.dates import format_month_year
from datetime import datetime
import re


class MissingDocxDependencyError(RuntimeError):
    pass


class ResumeDocxRenderer:
    """Create a simple, ATS-friendly DOCX without converting browser CSS."""

    def render(
        self,
        candidate: CandidateProfile,
        content: GeneratedResumeContent,
        *,
        target_title: str,
    ) -> bytes:
        resolved_title = target_title.strip()
        if not resolved_title:
            raise ValueError("target resume title must not be empty")
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches
        except ImportError as error:
            raise MissingDocxDependencyError(
                "install Word document support with: pip install -e ."
            ) from error

        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

        self._configure_styles(document)
        properties = document.core_properties
        properties.title = f"{candidate.full_name} - {resolved_title} Resume"
        properties.author = candidate.full_name
        properties.subject = "Tailored professional resume"

        name = document.add_paragraph(style="Title")
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name.add_run(candidate.full_name)

        contact_values = [candidate.email]
        if candidate.phone:
            contact_values.append(candidate.phone)
        if candidate.location:
            contact_values.append(candidate.location)
        contact_values.extend(
            url
            for url in (
                candidate.linkedin_url,
                candidate.github_url,
                candidate.website_url,
            )
            if url
        )
        contact = document.add_paragraph(" • ".join(contact_values), style="Subtitle")
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self._add_heading(document, "Professional Summary")
        summary = document.add_paragraph()
        summary.add_run(f"{resolved_title} — ")
        summary.add_run(content.professional_summary)

        if content.skills:
            self._add_heading(document, "Core Skills")
            document.add_paragraph(" • ".join(content.skills))

        if content.experience:
            self._add_heading(document, "Professional Experience")
            for role in content.experience:
                self._add_role(document, role)

        if content.career_highlights:
            self._add_heading(document, "Career Highlights")
            for value in content.career_highlights:
                paragraph = document.add_paragraph(style="List Bullet")
                if value.category:
                    paragraph.add_run(f"{value.category}: ")
                paragraph.add_run(value.description)

        if content.education:
            self._add_heading(document, "Education")
            for value in content.education:
                qualification = " in ".join(
                    part for part in (value.degree, value.field) if part
                )
                heading = qualification or value.institution
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.keep_with_next = True
                paragraph.add_run(heading)
                if qualification:
                    paragraph.add_run(f"\n{value.institution}")
                metadata = " • ".join(
                    part for part in (value.location, value.status) if part
                )
                if metadata:
                    paragraph.add_run(f"\n{metadata}").italic = True

        if content.certifications:
            self._add_heading(document, "Certifications")
            for value in content.certifications:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.keep_with_next = True
                paragraph.add_run(value.name)
                metadata = " • ".join(
                    part
                    for part in (
                        (
                            f"Issued {format_month_year(value.issued)}"
                            if value.issued
                            else None
                        ),
                        value.status,
                    )
                    if part
                )
                if metadata:
                    paragraph.add_run(f"\n{metadata}").italic = True

        output = BytesIO()
        document.save(output)
        return output.getvalue()

    @staticmethod
    def _configure_styles(document: Any) -> None:
        from docx.shared import Pt, RGBColor

        ink = RGBColor(23, 32, 51)
        muted = RGBColor(83, 97, 118)
        accent = RGBColor(29, 78, 137)

        normal = document.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(10)
        normal.font.bold = False
        normal.font.color.rgb = ink
        normal.paragraph_format.space_after = Pt(3)
        normal.paragraph_format.line_spacing = 1.08

        title = document.styles["Title"]
        title.font.name = "Arial"
        title.font.size = Pt(24)
        title.font.bold = False
        title.font.color.rgb = ink
        title.paragraph_format.space_after = Pt(3)

        subtitle = document.styles["Subtitle"]
        subtitle.font.name = "Arial"
        subtitle.font.size = Pt(9)
        subtitle.font.bold = False
        subtitle.font.italic = False
        subtitle.font.color.rgb = muted
        subtitle.paragraph_format.space_after = Pt(9)

        heading = document.styles["Heading 1"]
        heading.font.name = "Arial"
        heading.font.size = Pt(10.5)
        heading.font.bold = False
        heading.font.color.rgb = accent
        heading.paragraph_format.space_before = Pt(10)
        heading.paragraph_format.space_after = Pt(4)
        heading.paragraph_format.keep_with_next = True

        bullet = document.styles["List Bullet"]
        bullet.font.name = "Arial"
        bullet.font.size = Pt(10)
        bullet.font.bold = False
        bullet.paragraph_format.space_after = Pt(2)

    @staticmethod
    def _add_heading(document: Any, title: str) -> None:
        document.add_heading(title.upper(), level=1)

    @staticmethod
    def _add_role(document: Any, role: GeneratedResumeRole) -> None:
        from docx.enum.text import WD_TAB_ALIGNMENT
        from docx.shared import Emu, Pt, RGBColor

        heading = document.add_paragraph()
        heading.paragraph_format.keep_with_next = True
        heading.paragraph_format.space_before = Pt(4)
        heading.paragraph_format.space_after = Pt(1)
        section = document.sections[0]
        content_width = Emu(
            section.page_width - section.left_margin - section.right_margin
        )
        heading.paragraph_format.tab_stops.add_tab_stop(
            content_width,
            WD_TAB_ALIGNMENT.RIGHT,
        )
        heading.add_run(role.title).bold = True

        dates = ResumeDocxRenderer._dates(role)
        if dates:
            heading.add_run("\t")
            date_run = heading.add_run(dates)
            date_run.italic = True
            date_run.font.size = Pt(9)
            date_run.font.color.rgb = RGBColor(83, 97, 118)

        company = heading.add_run(f"\n{role.company}")
        company.bold = False
        company.font.color.rgb = RGBColor(83, 97, 118)
        if role.location:
            location = heading.add_run(f" • {role.location}")
            location.font.color.rgb = RGBColor(83, 97, 118)

        for responsibility in role.responsibilities:
            document.add_paragraph(responsibility, style="List Bullet")

    @staticmethod
    def _dates(role: GeneratedResumeRole) -> str:
        if role.start_date and role.end_date:
            start = datetime.strptime(role.start_date, "%Y-%m").strftime("%b %Y")
            end = datetime.strptime(role.end_date, "%Y-%m").strftime("%b %Y") if re.match("\d{4}-\d{2}", role.end_date) else "Current"
            return f"{start} – {end}"
        if role.start_date:
            return datetime.strptime(role.start_date, "%Y-%m").strftime("%b %Y")
        return datetime.strptime(role.end_date, "%Y-%m").strftime("%b %Y") if re.match("\d{4}-\d{2}", role.end_date) else "Current"
