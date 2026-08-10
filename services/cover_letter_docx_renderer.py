"""Render a concise cover letter as an editable one-page Word document."""

from __future__ import annotations

from datetime import date
from io import BytesIO
from typing import Any

from models import CandidateProfile, GeneratedCoverLetterContent, JobPosting

from .resume_docx_renderer import MissingDocxDependencyError


class CoverLetterDocxRenderer:
    def render(
        self,
        candidate: CandidateProfile,
        content: GeneratedCoverLetterContent,
        *,
        job: JobPosting,
        generated_on: date | None = None,
    ) -> bytes:
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches
        except ImportError as error:
            raise MissingDocxDependencyError(
                "install Word document support with: pip install -e ."
            ) from error

        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.78)
        section.right_margin = Inches(0.78)
        self._configure_styles(document)

        properties = document.core_properties
        properties.title = f"{candidate.full_name} - {job.company} Cover Letter"
        properties.author = candidate.full_name
        properties.subject = f"Application for {job.title}"

        name = document.add_paragraph(style="Title")
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name.add_run(candidate.full_name)

        contact_values = [candidate.email]
        contact_values.extend(
            value
            for value in (
                candidate.phone,
                candidate.location,
                candidate.linkedin_url,
                candidate.website_url,
            )
            if value
        )
        contact = document.add_paragraph(" • ".join(contact_values), style="Subtitle")
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

        resolved_date = generated_on or date.today()
        document.add_paragraph(_format_date(resolved_date))
        recipient = document.add_paragraph()
        recipient.add_run(job.company)
        recipient.add_run(f"\n{job.title}")
        if job.location:
            recipient.add_run(f"\n{job.location}")

        document.add_paragraph(f"Dear {job.company} Hiring Team,")
        for paragraph in content.paragraphs:
            document.add_paragraph(paragraph.text)
        document.add_paragraph(f"Sincerely,\n{candidate.full_name}")

        output = BytesIO()
        document.save(output)
        return output.getvalue()

    @staticmethod
    def _configure_styles(document: Any) -> None:
        from docx.shared import Pt, RGBColor

        ink = RGBColor(23, 32, 51)
        muted = RGBColor(83, 97, 118)

        normal = document.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(10.5)
        normal.font.color.rgb = ink
        normal.paragraph_format.space_after = Pt(8)
        normal.paragraph_format.line_spacing = 1.08

        title = document.styles["Title"]
        title.font.name = "Arial"
        title.font.size = Pt(22)
        title.font.bold = False
        title.font.color.rgb = ink
        title.paragraph_format.space_after = Pt(2)

        subtitle = document.styles["Subtitle"]
        subtitle.font.name = "Arial"
        subtitle.font.size = Pt(9)
        subtitle.font.bold = False
        subtitle.font.italic = False
        subtitle.font.color.rgb = muted
        subtitle.paragraph_format.space_after = Pt(14)


def _format_date(value: date) -> str:
    return f"{value.strftime('%B')} {value.day}, {value.year}"
